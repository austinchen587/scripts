# document_processor/image_extractor.py
import os
import zipfile
import tempfile
from pathlib import Path
from typing import List
from pdf2image import convert_from_path
from .ocr_processor import extract_from_image_with_ollama, extract_text_from_image_ocr
from .config import POPPLER_BIN_PATH, DOWNLOAD_DIR
import pdfplumber
import time  # <--- 新增这行

class OfficeImageExtractor:
    """提取Office文档中的嵌入图片并进行文字识别"""
    
    def __init__(self, debug: bool = True):
        self.debug = debug
        self.supported_extensions = ['.docx', '.xlsx', '.pptx', '.doc', '.xls', '.ppt', '.pdf']

        # --- 放大超时限制，确保4060显卡有足够时间处理多张图片 (20分钟) ---
        self.MAX_IMAGE_PROCESS_TIME = 1200
    
    def extract_from_docx(self, file_path: Path, project_name: str = "") -> str:
        """从Word文档中提取文字和图片文字"""
        try:
            if self.debug:
                print(f"[IMAGE-EXTRACT] 🖼️ 处理Word文档: {file_path.name}")
            
            # 1. 使用python-docx提取文字
            text_content = self._extract_docx_text(file_path)
            
            # 2. 提取嵌入的图片并识别
            image_texts = self._extract_docx_images(file_path, project_name)
            
            # 合并结果
            if image_texts:
                text_content += "\n\n--- 图片提取内容 ---\n" + "\n".join(image_texts)
            
            return text_content
            
        except Exception as e:
            print(f"[IMAGE-EXTRACT] ⚠️ 提取Word文档失败: {e}")
            return ""
    
    def extract_from_xlsx(self, file_path: Path, project_name: str = "") -> str:
        """从Excel文档中提取文字和图片文字"""
        try:
            if self.debug:
                print(f"[IMAGE-EXTRACT] 📊 处理Excel文档: {file_path.name}")
            
            # 1. 使用openpyxl提取单元格文字
            text_content = self._extract_xlsx_text(file_path)
            
            # 2. 尝试提取嵌入图片（Excel中的图片处理较复杂）
            image_texts = self._extract_xlsx_images(file_path, project_name)
            
            if image_texts:
                text_content += "\n\n--- 图片提取内容 ---\n" + "\n".join(image_texts)
            
            return text_content
            
        except Exception as e:
            print(f"[IMAGE-EXTRACT] ⚠️ 提取Excel文档失败: {e}")
            return ""
        
    def extract_from_pdf(self, file_path: Path, project_name: str = "") -> str:
        """从PDF文档中提取文字和图片文字"""
        try:
            if self.debug:
                print(f"[IMAGE-EXTRACT] 📄 处理PDF文档: {file_path.name}")
            
            # 1. 首先尝试原生文字提取
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text_content = ""
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content += text + "\n"
                
                if text_content and len(text_content) > 50:
                    if self.debug:
                        print(f"[IMAGE-EXTRACT] ✅ 获取原生PDF文字: {len(text_content)}字符")
            except Exception as e:
                if self.debug:
                    print(f"[IMAGE-EXTRACT] ⚠️ 原生PDF提取失败: {e}")
                text_content = ""
            
            # 2. 提取PDF中的图片并识别
            image_texts = self._extract_pdf_images(file_path, project_name)
            
            # 合并结果
            if image_texts:
                if text_content:
                    text_content += "\n\n--- 图片提取内容 ---\n"
                else:
                    text_content = "--- 图片提取内容 ---\n"
                text_content += "\n".join(image_texts)
            
            # 3. 如果还是没有内容，直接返回空
            if not text_content or len(text_content.strip()) < 10:
                return ""
                
            return text_content
            
        except Exception as e:
            print(f"[IMAGE-EXTRACT] 💥 提取PDF文档失败: {e}")
            return ""
    
    def _extract_pdf_images(self, file_path: Path, project_name: str) -> List[str]:
        """从PDF中提取图片"""
        try:
            if self.debug:
                print(f"[IMAGE-EXTRACT] 🔧 开始提取PDF图片: {file_path.name}")
            
            # 验证文件存在
            if not file_path.exists():
                print(f"[IMAGE-EXTRACT] ❌ PDF文件不存在: {file_path}")
                return []
            
            # 临时存储所有图片路径
            temp_file_paths = []
            image_texts = []
            
            try:
                # 获取Poppler路径
                poppler_path = POPPLER_BIN_PATH
                
                if poppler_path:
                    # 检查Poppler路径有效性
                    if not os.path.exists(poppler_path):
                        print(f"[IMAGE-EXTRACT] ⚠️ Poppler路径不存在: {poppler_path}")
                        poppler_path = None
                    else:
                        # 检查关键可执行文件
                        if os.name == 'nt':  # Windows
                            pdftoppm_exe = os.path.join(poppler_path, 'pdftoppm.exe')
                            if not os.path.exists(pdftoppm_exe):
                                print(f"[IMAGE-EXTRACT] ⚠️ 找不到pdftoppm.exe: {pdftoppm_exe}")
                                poppler_path = None
                
                if self.debug:
                    print(f"[IMAGE-EXTRACT] 📄 正在转换PDF: {file_path.name}...")
                    print(f"[IMAGE-EXTRACT] 🔧 Poppler路径: {poppler_path or '自动检测'}")
                
                # 将PDF转换为图片
                images = convert_from_path(
                    str(file_path),
                    dpi=150,                    # 分辨率
                    fmt="png",                   # 输出格式
                    poppler_path=poppler_path,   # Poppler路径
                    thread_count=1,              # 单线程避免问题
                    use_pdftocairo=True          # 使用pdftocairo可能更稳定
                )
                
                if not images:
                    print(f"[IMAGE-EXTRACT] ⚠️ PDF转换后无图片输出")
                    return []
                
                if self.debug:
                    print(f"[IMAGE-EXTRACT] ✅ PDF转换成功: {len(images)} 页")

                # --- 新增：记录开始时间 ---
                start_time = time.time()
                
                for i, img in enumerate(images):
                    # --- 新增：循环内检查超时 ---
                    if time.time() - start_time > self.MAX_IMAGE_PROCESS_TIME:
                        print(f"[IMAGE-EXTRACT] 🛑 PDF图片提取超时 (> {self.MAX_IMAGE_PROCESS_TIME}s)，停止后续页面")
                        image_texts.append("\n[警告] 图片过多，后续页面提取已截断...")
                        break
                    # -------------------------

                    try:
                        # 创建绝对安全的临时文件路径
                        import uuid
                        temp_dir = tempfile.gettempdir()
                        unique_id = uuid.uuid4().hex[:8]
                        temp_filename = f"pdf_extract_{os.getpid()}_{unique_id}_page_{i+1}.png"
                        temp_img_path = Path(temp_dir) / temp_filename
                        
                        # 保存图片
                        img.save(temp_img_path, "PNG", quality=85)
                        temp_file_paths.append(temp_img_path)
                        
                        if self.debug:
                            print(f"[IMAGE-EXTRACT] 💾 保存临时图片: {temp_img_path.name}")
                        
                        # 验证文件保存成功
                        if not temp_img_path.exists():
                            print(f"[IMAGE-EXTRACT] ⚠️ 临时文件未保存成功: {temp_img_path}")
                            continue
                        
                        # 识别图片文字
                        img_text = self._process_image_for_text(temp_img_path, project_name, f"page_{i+1}")
                        if img_text and len(img_text.strip()) > 10:
                            image_texts.append(f"[PDF第{i+1}页图片]\n{img_text}")
                            if self.debug:
                                print(f"[IMAGE-EXTRACT] ✅ 第{i+1}页提取成功: {len(img_text)}字符")
                        
                    except Exception as img_e:
                        if self.debug:
                            print(f"[IMAGE-EXTRACT] ⚠️ 处理PDF第{i+1}页失败: {img_e}")
                        continue
                
            except Exception as conv_e:
                print(f"[IMAGE-EXTRACT] ❌ PDF转换失败: {conv_e}")
                import traceback
                traceback.print_exc()
                return []
            finally:
                # 清理所有临时文件
                failed_cleanups = 0
                for temp_path in temp_file_paths:
                    try:
                        if temp_path and temp_path.exists():
                            temp_path.unlink(missing_ok=True)
                            if self.debug:
                                print(f"[IMAGE-EXTRACT] 🗑️ 清理临时文件: {temp_path.name}")
                    except Exception as cleanup_e:
                        failed_cleanups += 1
                        if self.debug:
                            print(f"[IMAGE-EXTRACT] ⚠️ 清理临时文件失败: {cleanup_e}")
                
                if failed_cleanups > 0:
                    print(f"[IMAGE-EXTRACT] ⚠️ {failed_cleanups} 个临时文件清理失败")
            
            if self.debug:
                print(f"[IMAGE-EXTRACT] ✅ 提取完成: 共 {len(image_texts)} 页有内容")
            
            return image_texts
            
        except Exception as e:
            print(f"[IMAGE-EXTRACT] 💥 提取PDF图片失败: {e}")
            import traceback
            traceback.print_exc()
            return []

    
    def _extract_docx_text(self, file_path: Path) -> str:
        """使用python-docx提取Word文字"""
        try:
            import docx
            doc = docx.Document(file_path)
            
            # 提取段落文字
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # 提取表格文字
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n".join(paragraphs)
            
        except ImportError:
            print("[IMAGE-EXTRACT] ⚠️ python-docx未安装，无法提取Word文字")
            return ""
        except Exception as e:
            print(f"[IMAGE-EXTRACT] ⚠️ 提取Word文字失败: {e}")
            return ""
    
    def _extract_docx_images(self, file_path: Path, project_name: str) -> List[str]:
        """从Word文档中提取图片并进行OCR"""
        image_texts = []
        
        try:
            # Word文档本质上是zip文件
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # 查找图片文件
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/') and 
                              f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp'))]
                
                if self.debug:
                    print(f"[IMAGE-EXTRACT] 🔍 发现 {len(image_files)} 张图片")


                # --- 新增：记录开始时间 ---
                start_time = time.time()
                
                for img_file in image_files:

                    # --- 新增：循环内检查超时 ---
                    if time.time() - start_time > self.MAX_IMAGE_PROCESS_TIME:
                        print(f"[IMAGE-EXTRACT] 🛑 Word图片提取超时 (> {self.MAX_IMAGE_PROCESS_TIME}s)，停止提取剩余图片")
                        image_texts.append("\n[警告] 图片过多，提取已截断...")
                        break
                    # -------------------------



                    try:
                        # 提取图片数据
                        img_data = docx_zip.read(img_file)
                        
                        # 保存为临时文件
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                            tmp.write(img_data)
                            temp_path = Path(tmp.name)
                        
                        # 识别图片中的文字
                        img_text = self._process_image_for_text(temp_path, project_name, img_file)
                        if img_text:
                            image_texts.append(f"[图片: {os.path.basename(img_file)}]\n{img_text}")
                        
                        # 清理临时文件
                        temp_path.unlink(missing_ok=True)
                        
                    except Exception as img_e:
                        if self.debug:
                            print(f"[IMAGE-EXTRACT] ⚠️ 处理图片失败 {img_file}: {img_e}")
                        continue
            
            return image_texts
            
        except Exception as e:
            print(f"[IMAGE-EXTRACT] ⚠️ 提取Word图片失败: {e}")
            return []
    
    def _extract_xlsx_text(self, file_path: Path) -> str:
        """【关键修复】使用pandas提取Excel文字，自动去除幽灵空行"""
        try:
            import pandas as pd
            # 读取所有sheet
            dfs = pd.read_excel(file_path, sheet_name=None, header=None)
            text_content = ""
            
            for sheet_name, df in dfs.items():
                # 删除全为空的行和列
                df = df.dropna(how='all')
                df = df.dropna(axis=1, how='all')
                
                if not df.empty:
                    # 限制数据量，防止极个别超大文件
                    if len(df) > 5000:
                        df = df.head(5000)
                        text_content += f"工作表: {sheet_name} (仅显示前5000行)\n"
                    else:
                        text_content += f"工作表: {sheet_name}\n"
                    
                    # 转换为CSV文本
                    sheet_text = df.to_csv(index=False, header=False, sep='\t')
                    text_content += sheet_text + "\n\n"
            
            return text_content.strip()
            
        except ImportError:
            print("[IMAGE-EXTRACT] ⚠️ pandas未安装")
            return ""
        except Exception as e:
            print(f"[IMAGE-EXTRACT] ⚠️ 提取Excel文字失败: {e}")
            return ""
    
    def _extract_xlsx_images(self, file_path: Path, project_name: str) -> List[str]:
        """从Excel中提取图片（Excel中的图片处理较复杂）"""
        if self.debug:
            print("[IMAGE-EXTRACT] ℹ️ Excel图片提取暂未实现")
        return []
    
    def _extract_pdf_images(self, file_path: Path, project_name: str) -> List[str]:
        """从PDF中提取图片"""
        try:
            if self.debug:
                print(f"[IMAGE-EXTRACT] 🔧 开始提取PDF图片: {file_path.name}")
            
            # 【修复】将DOWNLOAD_DIR转换为Path对象
            download_dir = Path(DOWNLOAD_DIR)
            
            # 验证文件存在
            if not file_path.exists():
                print(f"[IMAGE-EXTRACT] ❌ PDF文件不存在: {file_path}")
                return []
            
            # 创建项目特定的临时目录
            import uuid
            temp_dir = download_dir / f"temp_{uuid.uuid4().hex[:8]}"
            temp_dir.mkdir(parents=True, exist_ok=True)
            
            temp_file_paths = []
            image_texts = []
            
            try:
                # 获取Poppler路径
                poppler_path = POPPLER_BIN_PATH
                
                if self.debug:
                    print(f"[IMAGE-EXTRACT] 📄 正在转换PDF: {file_path.name}...")
                    print(f"[IMAGE-EXTRACT] 🔧 Poppler路径: {poppler_path or '自动检测'}")
                    print(f"[IMAGE-EXTRACT] 📁 临时目录: {temp_dir}")
                
                # 将PDF转换为图片
                if not poppler_path or not os.path.exists(poppler_path):
                    # 如果未设置Poppler路径，让系统自动查找
                    images = convert_from_path(str(file_path), dpi=150, fmt="png")
                else:
                    images = convert_from_path(
                        str(file_path),
                        dpi=150,
                        fmt="png",
                        poppler_path=poppler_path
                    )
                
                if not images:
                    print(f"[IMAGE-EXTRACT] ⚠️ PDF转换后无图片输出")
                    return []
                
                if self.debug:
                    print(f"[IMAGE-EXTRACT] ✅ PDF转换成功: {len(images)} 页")
                
                for i, img in enumerate(images):
                    try:
                        # 【修复】使用项目临时目录
                        temp_img_path = temp_dir / f"page_{i+1}.png"
                        
                        # 保存图片
                        img.save(temp_img_path, "PNG", quality=85)
                        temp_file_paths.append(temp_img_path)
                        
                        if self.debug:
                            print(f"[IMAGE-EXTRACT] 💾 保存图片: {temp_img_path}")
                        
                        # 验证文件保存成功
                        if not temp_img_path.exists():
                            print(f"[IMAGE-EXTRACT] ⚠️ 图片未保存成功: {temp_img_path}")
                            continue
                        
                        # 识别图片文字
                        img_text = self._process_image_for_text(temp_img_path, project_name, f"page_{i+1}")
                        if img_text and len(img_text.strip()) > 10:
                            image_texts.append(f"[PDF第{i+1}页]\n{img_text}")
                            if self.debug:
                                print(f"[IMAGE-EXTRACT] ✅ 第{i+1}页提取成功: {len(img_text)}字符")
                        
                    except Exception as img_e:
                        if self.debug:
                            print(f"[IMAGE-EXTRACT] ⚠️ 处理PDF第{i+1}页失败: {img_e}")
                        continue
                        
            except Exception as conv_e:
                print(f"[IMAGE-EXTRACT] ❌ PDF转换失败: {conv_e}")
                import traceback
                traceback.print_exc()
                return []
            finally:
                # 清理所有临时文件
                for temp_path in temp_file_paths:
                    try:
                        if temp_path.exists():
                            temp_path.unlink()
                    except Exception as cleanup_e:
                        if self.debug:
                            print(f"[IMAGE-EXTRACT] ⚠️ 清理图片文件失败: {cleanup_e}")
                
                # 删除临时目录
                try:
                    temp_dir.rmdir()
                    if self.debug:
                        print(f"[IMAGE-EXTRACT] 🗑️ 清理临时目录: {temp_dir}")
                except Exception as dir_e:
                    if self.debug:
                        print(f"[IMAGE-EXTRACT] ⚠️ 清理临时目录失败: {dir_e}")
            
            if self.debug:
                print(f"[IMAGE-EXTRACT] ✅ 提取完成: 共 {len(image_texts)} 页有内容")
            
            return image_texts
            
        except Exception as e:
            print(f"[IMAGE-EXTRACT] 💥 提取PDF图片失败: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    def _process_image_for_text(self, image_path: Path, project_name: str, image_source: str) -> str:
        """处理单张图片，提取文字"""
        try:
            if self.debug:
                print(f"[IMAGE-EXTRACT] 🔍 处理图片: {image_source}")
            
            # 优先使用Ollama视觉模型
            try:
                ollama_text = extract_from_image_with_ollama(image_path, project_name)
                if ollama_text and len(ollama_text.strip()) > 10:
                    return ollama_text
            except Exception as ollama_e:
                if self.debug:
                    print(f"[IMAGE-EXTRACT] ⚠️ Ollama-VL失败: {ollama_e}")
            
            # 备用方案：OCR
            try:
                ocr_text = extract_text_from_image_ocr(image_path)
                if ocr_text and len(ocr_text.strip()) > 10:
                    return ocr_text
            except Exception as ocr_e:
                if self.debug:
                    print(f"[IMAGE-EXTRACT] ⚠️ OCR失败: {ocr_e}")
            
            return ""
            
        except Exception as e:
            print(f"[IMAGE-EXTRACT] ⚠️ 处理图片文字失败 {image_source}: {e}")
            return ""
    
    def extract_with_images(self, file_path: str, project_name: str = "") -> str:
        """智能提取文档文字，包含嵌入图片识别"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == '.docx':
            return self.extract_from_docx(path, project_name)
        elif suffix in ['.xlsx', '.xls']:
            return self.extract_from_xlsx(path, project_name)
        elif suffix == '.pdf':
            return self.extract_from_pdf(path, project_name)
        elif suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            # 直接处理图片文件
            return self._process_image_for_text(path, project_name, "direct_image")
        else:
            print(f"[IMAGE-EXTRACT] ⚠️ 不支持的文件类型: {suffix}")
            return ""

# 创建全局实例
image_extractor = OfficeImageExtractor(debug=True)

def extract_with_embedded_images(file_path: str, project_name: str = "") -> str:
    """提取文档文字，包括嵌入图片内容"""
    return image_extractor.extract_with_images(file_path, project_name)
