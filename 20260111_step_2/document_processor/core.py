# document_processor/core.py
import os
from pathlib import Path
from typing import Optional
from .image_extractor import extract_with_embedded_images
from .libreoffice_utils import office_to_pdf
from .pdf_processor import extract_text_native, pdf_to_images
from .ocr_processor import extract_from_image_with_ollama, extract_text_from_image_ocr

# document_processor/core.py - 关键部分修复
def extract_text_enhanced(file_path: str, project_name: str = "") -> str:
    """
    增强的文档提取：支持嵌入图片识别
    策略：优先提取原生文字，然后尝试提取图片文字
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    if not project_name:
        project_name = file_path.stem
    print(f"[DOC] 📄 处理文件: {file_path.name} (类型: {suffix})")
    
    # === 直接处理图片文件 ===
    if suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
        print(f"[DOC] 🖼️ 检测到图片文件，启动视觉分析...")
        # 1. 尝试视觉模型
        text = extract_from_image_with_ollama(file_path, project_name)
        # 2. 抢救逻辑
        if not text or len(text.strip()) < 5:
            print(f"[DOC] ⚠️ 视觉模型无响应或内容过少，正在启动 Tesseract OCR 抢救...")
            text = extract_text_from_image_ocr(file_path)
        return text
    
    # === 处理Office文档和PDF ===
    # 第一步：尝试原生文字提取（快速）
    native_text = ""
    
    if suffix == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            
            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 对于特别长的段落，适当添加换行
                    if len(text) > 100:
                        # 添加一些标点处的换行
                        text = text.replace('。', '。\n').replace('；', '；\n')
                    paragraphs.append(text)
            
            # 提取表格内容 - 确保完整提取
            for table in doc.tables:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        # 处理技术参数的长文本
                        if cell_text:
                            # 简单的格式化：用空格分隔参数
                            cell_text = cell_text.replace(':', ': ').replace(';', '; ')
                        if cell_text:
                            row_cells.append(cell_text)
                    
                    if row_cells and any(cell for cell in row_cells):
                        # 使用制表符分隔单元格，便于识别表格结构
                        paragraph_text = "\t".join(row_cells)
                        if paragraph_text.strip():
                            paragraphs.append(paragraph_text)
            
            native_text = "\n".join(paragraphs)
            
        except ImportError:
            print("[DOC] ⚠️ python-docx未安装，跳过原生提取")
        except Exception as e:
            print(f"[DOC] ⚠️ 提取Word文字失败: {e}")
    
    elif suffix in ['.doc', '.wps']:
        try:
            # .doc或.wps文件转换为PDF再处理
            print(f"[DOC] 🔄 {suffix}文件需要转换为PDF")
            pdf_path = office_to_pdf(file_path)
            from .pdf_processor import extract_text_native
            native_text = extract_text_native(pdf_path)
            # 清理临时PDF文件
            try:
                pdf_path.unlink(missing_ok=True)
            except:
                pass
        except Exception as e:
            print(f"[DOC] ⚠️ 提取.doc文件失败: {e}")
    
    elif suffix in ['.xlsx', '.et']:
        # 【关键修复】使用pandas读取，彻底解决幽灵空行问题
        try:
            import pandas as pd
            # 读取所有sheet，header=None防止表头被漏掉
            dfs = pd.read_excel(file_path, sheet_name=None, header=None)
            sheet_texts = []
            for sheet_name, df in dfs.items():
                # 1. 删除全为空的行
                df = df.dropna(how='all')
                # 2. 删除全为空的列
                df = df.dropna(axis=1, how='all')
                
                if not df.empty:
                    # 转换为制表符分隔的文本，保留表格结构
                    text = df.to_csv(index=False, header=False, sep='\t')
                    # 限制单个sheet长度，防止极端情况
                    if len(text) > 50000:
                        text = text[:50000] + "\n...[内容过长已截断]..."
                    sheet_texts.append(f"工作表: {sheet_name}\n{text}")
            
            native_text = "\n\n".join(sheet_texts)
            
        except ImportError:
            print("[DOC] ⚠️ pandas未安装，无法提取Excel文字")
        except Exception as e:
            print(f"[DOC] ⚠️ 提取Excel文字失败: {e}")
    
    elif suffix == '.xls':  # 添加.xls文件处理
        try:
            print(f"[DOC] 🔄 .xls文件需要转换格式或使用xlrd")
            # 方案1: 转换为PDF
            pdf_path = office_to_pdf(file_path)
            from .pdf_processor import extract_text_native
            native_text = extract_text_native(pdf_path)
            # 清理临时PDF文件
            try:
                pdf_path.unlink(missing_ok=True)
            except:
                pass
        except Exception as e:
            print(f"[DOC] ⚠️ 转换.xls失败: {e}")
            # 方案2: 尝试使用xlrd（如果安装）
            try:
                import xlrd
                workbook = xlrd.open_workbook(str(file_path))
                sheet_texts = []
                for sheet_name in workbook.sheet_names():
                    sheet = workbook.sheet_by_name(sheet_name)
                    rows = []
                    for row_idx in range(sheet.nrows):
                        row_cells = [str(sheet.cell_value(row_idx, col_idx)) 
                                   for col_idx in range(sheet.ncols)]
                        if any(cell.strip() for cell in row_cells):
                            rows.append("\t".join(row_cells))
                    if rows:
                        sheet_texts.append(f"工作表: {sheet_name}\n" + "\n".join(rows))
                native_text = "\n\n".join(sheet_texts)
            except ImportError:
                print("[DOC] ⚠️ xlrd未安装，无法读取.xls文件")
            except Exception as e2:
                print(f"[DOC] ⚠️ 使用xlrd读取.xls失败: {e2}")
    
    elif suffix == '.pdf':
        try:
            from .pdf_processor import extract_text_native
            native_text = extract_text_native(file_path)
        except Exception as e:
            print(f"[DOC] ⚠️ 提取PDF文字失败: {e}")
    
    # 检查原生文字是否足够
    if native_text and len(native_text) > 50:  # 降低阈值
        print(f"[DOC] ✅ 原生文字提取成功: {len(native_text)} 字符")
        
        # 对于不支持图片提取的文件类型，跳过图片提取
        unsupported_for_images = ['.doc', '.xls', '.pdf', '.wps', '.et']
        if suffix not in unsupported_for_images:
            try:
                image_text = extract_with_embedded_images(file_path, project_name)
                if image_text and len(image_text) > 30:
                    print(f"[DOC] 🖼️ 从嵌入图片中提取了额外内容")
                    native_text += "\n\n--- 图片提取内容 ---\n" + image_text
            except Exception as e:
                print(f"[DOC] ⚠️ 图片提取失败: {e}")
        
        return native_text
    
    # 第二步：原生文字不足，使用增强提取
    print(f"[DOC] 🔍 原生文字不足或提取失败，使用增强提取")
    
    try:
        # 对于某些文件类型，直接返回原生文本（即使很少）
        if native_text:
            print(f"[DOC] 🔄 使用已提取的原生文字（{len(native_text)} 字符）")
            return native_text
        
        # 尝试使用图片提取模块
        unsupported_types = ['.doc', '.xls', '.wps', '.et']
        if suffix not in unsupported_types:
            full_text = extract_with_embedded_images(file_path, project_name)
            if full_text and len(full_text) > 50:
                print(f"[DOC] ✅ 增强提取成功: {len(full_text)} 字符")
                return full_text
    except Exception as e:
        print(f"[DOC] ⚠️ 增强提取失败: {e}")
    
    # 第三步：所有方法都失败
    print(f"[DOC] ❌ 所有提取方法均失败，返回空文本")
    return ""


def extract_text(file_path: str, project_name: str = "") -> str:
    """
    智能提取文档文本：
    - 图片文件：直接OCR或Ollama-VL处理
    - Office文件 → PDF → 原生文字 or OCR
    - PDF文件 → 原生文字 or OCR
    
    注意：这是原始函数，建议使用 extract_text_enhanced 以获得更好的图片处理能力
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    # 如果没有传入项目名，使用文件名
    if not project_name:
        project_name = file_path.stem

    print(f"[DOC] 📄 处理文件: {file_path.name} (类型: {suffix})")

    # === 直接处理图片文件 ===
    if suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
        print(f"[DOC] 🖼️ 检测到图片文件，使用OCR处理")
        # 首先尝试Ollama-VL
        vl_text = extract_from_image_with_ollama(file_path, project_name)
        if not vl_text or len(vl_text.strip()) < 5:
            return extract_text_from_image_ocr(file_path)
        return vl_text

    # === 处理文档文件 ===
    # Step 1: 统一转为 PDF（如果是 Office）
    if suffix in ['.doc', '.docx', '.xls', '.xlsx', '.wps', '.et']:
        try:
            pdf_path = office_to_pdf(file_path)
        except Exception as e:
            print(f"[DOC] ⚠️ LibreOffice 转换失败: {e}")
            return ""
    elif suffix == '.pdf':
        pdf_path = file_path
    else:
        print(f"[DOC] ⚠️ 不支持的文件类型: {suffix}")
        return ""

    # Step 2: 尝试原生文字提取
    text = extract_text_native(pdf_path)
    if text.strip():
        print(f"[DOC] ✅ 原生文本提取成功: {len(text)} 字符")
        return text

    # Step 3: OCR 流程（扫描件）
    print(f"[DOC] 🖼️ 检测到扫描版 PDF，启动 OCR")
    try:
        img_paths = pdf_to_images(pdf_path)
        if not img_paths:
            return ""
            
        results = []
        for img_path in img_paths:
            # 优先Ollama-VL
            vl_text = extract_from_image_with_ollama(img_path, project_name)
            if vl_text:
                results.append(vl_text)
            else:
                # 备用OCR
                ocr_text = extract_text_from_image_ocr(img_path)
                if ocr_text:
                    results.append(ocr_text)
            
            # 清理临时图片文件
            try:
                img_path.unlink()
            except:
                pass
                
        # 清理临时目录
        try:
            img_path.parent.rmdir()
        except:
            pass
            
        return "\n".join(results)
    except Exception as e:
        print(f"[DOC] 💥 OCR 全流程失败: {e}")
        return ""
